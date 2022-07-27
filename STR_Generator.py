# Filename: word editing with python.py
# Author:   David Preiss
# Date:     7/19/2022
# Edited:   7/21/2022

###### TABLE OF CONTENTS:

    # OVERVIEW
    # IMPORTS
    # FUNCTIONS
    # CODE INTRO
    # CODE START

###### OVERVIEW:

# This python script is supposed to automate the completion of STR documents
# Run this script in a folder with exactly 1 .docx file and exactly 1 .csv file to skip the menu
# Run under normal conditions to gain access to the menu

###### IMPORTS:


import os                   # for changing working directory
try:
    from docx import Document   # for opening/editing/saving docx
except:
    os.system("pip install python-docx")
    from docx import Document
import datetime             # for naming the folder
from datetime import date   # for putting the date of the signature
import pathlib              # in case we need the path
import csv                  # for creating/reading from the csv

os.system("echo Hello from the other side!")
###### FUNCTIONS:

def AutoFiller_For_C_and_D(sourceFile, destinationFile, inputTest, inputResults, inputTesterName, inputDate, inputAnomalies):
    # This function accepts a 2 filepaths for an STR .docx file and the information for sections C and D,
    # it then uses the first file as a template, writes the information, and saves to the second path

    # Open the template file
    document = Document(sourceFile)

    # Determine if this .docx file was created by this program
    if (document.core_properties.title == 'Generated by David Preiss'):
        # if it was, edit the corresponding textboxes with the appropriate information
        document.tables[3].rows[0].cells[0].text = inputTest
        document.tables[4].rows[0].cells[0].text = inputResults
        document.tables[5].rows[1].cells[0].text = inputTesterName
        document.tables[5].rows[1].cells[1].text = inputDate
        document.tables[5].rows[1].cells[3].text = str(inputAnomalies)
    else:
        # if it wasn't, assume format matches siemens official STR docx
        document.tables[0].rows[14].cells[0].text = inputTest
        document.tables[0].rows[18].cells[0].text = inputResults
        document.tables[0].rows[21].cells[0].text = inputTesterName
        document.tables[0].rows[21].cells[5].text = inputDate
        document.tables[0].rows[21].cells[16].text = str(inputAnomalies)

    # save the edited STR to the chosen destination path
    document.save(destinationFile)
    return

def STR_Generation_Script():

    # Defunct: This function is not called

    s_targetDoc = "Engineering Test Report Form-B2B.docx"

    s_readDoc = s_targetDoc
    s_FolderName = str(datetime.datetime.now()).replace(":","_")

    print(s_FolderName)
    os.mkdir(s_FolderName)

    s_Name_of_CSV = "B2B STPr csv.csv"
    with open(s_Name_of_CSV, newline='') as csvfile:
        spamreader = csv.reader(csvfile, delimiter=',', quotechar='|')
        for row in spamreader:

            # print(row[3])

            if (row[1]!="Owner" and row[3]!=""):
                b_newDoc = False
                print("Test: ",row[0],"\t Date: ",row[3],'\n name: ',row[1])

                #Open a STR docx with A and B filled in, but not C or D
                # Fill in C with the Test Number
                # Fill in D with Results, Name, Current Date, and anomolies
                # save as a new Document under a newfolder, named for the test
                AutoFiller_For_C_and_D(s_readDoc, s_FolderName+'/'+row[0]+'.docx', row[0], 'Test Passed', row[1], row[3], 0)

def GenCSV_template():
    # This function Generates a template.csv file for STPrs
    print("Generating template.csv file...")

    try:
        # open the file in the write mode
        f = open('template.csv', 'w')

        # create the csv writer
        writer = csv.writer(f)

        # Create header
        header = ['Test Number', 'Tester Signature', 'Date Started', 'Date Completed', 'Result Pass/Fail', 'Anomalies', 'Additional Notes']

        # write header to the csv file
        writer.writerow(header)

        print('Successfully generated template.csv')
    except:
        print('ERROR! Failed to generate csv file')
    return

def GenDOCX_template():
    # This function Generates a template.docx file for STRs
    print("Generating template.docx file...")

    try:

        newDoc = Document() # creates a new blank .docx file
        newDoc.core_properties.title = 'Generated by David Preiss'

        # a. Identification:
        newDoc.add_paragraph('a. Identification:')

        table1 = newDoc.add_table(rows=2,cols=4)
        table1.style = 'Table Grid'
        table1.style.font.size = 100000
        table1.rows[0].cells[0].text = 'Date:'
        table1.rows[0].cells[1].text = 'Project Number:'
        table1.rows[0].cells[2].text = 'Project Name:'
        table1.rows[0].cells[3].text = 'Report Number:'

        table2 = newDoc.add_table(rows=2,cols=3)
        table2.style = 'Table Grid'
        table2.rows[0].cells[0].text = 'Test Procedure-Document Number:'
        table2.rows[0].cells[1].text = 'Rev:'
        table2.rows[0].cells[2].text = 'Document Title:'

        newDoc.add_paragraph('b. Type of Test:')

        table3 = newDoc.add_table(rows=5,cols=8)
        table3.style = 'Table Grid'
        table3.rows[0].cells[1].text = 'Software:'
        table3.rows[0].cells[3].text = 'System:'
        table3.rows[0].cells[5].text = 'Qualification:'
        table3.rows[0].cells[7].text = 'Field / Acceptance:'

        table3.rows[1].cells[1].text = 'Integration Test'
        table3.rows[1].cells[3].text = 'Integration Test'
        table3.rows[1].cells[5].text = 'Functional Test'
        table3.rows[1].cells[7].text = 'FAI / Acceptance'

        table3.rows[2].cells[1].text = 'Requirements Test'
        table3.rows[2].cells[3].text = 'Functional Test'
        table3.rows[2].cells[5].text = 'Temperature Test'
        table3.rows[2].cells[7].text = 'Static Field Test'
        
        table3.rows[3].cells[5].text = 'Shock & Vibration'
        table3.rows[3].cells[7].text = 'Dynamic Field Test'
        
        table3.rows[4].cells[5].text = 'EMI / EMC Test'

        newDoc.add_paragraph('c. Test objective(s):')

        table4 = newDoc.add_table(rows=1,cols=1)
        table4.style = 'Table Grid'

        newDoc.add_paragraph('d. Summary of Test Results:')

        table5 = newDoc.add_table(rows=1,cols=1)
        table5.style = 'Table Grid'

        table6 = newDoc.add_table(rows=2,cols=4)
        table6.style = 'Table Grid'
        table6.rows[0].cells[0].text = 'Test Conducted by:'
        table6.rows[0].cells[1].text = 'Date:'
        table6.rows[0].cells[3].text = 'Number of Anomalies Reported:'

        newDoc.add_paragraph('e. Final disposition/audit of test result(s):')

        table7 = newDoc.add_table(rows=1,cols=1)
        table7.style = 'Table Grid'

        table8 = newDoc.add_table(rows=2,cols=4)
        table8.style = 'Table Grid'
        table8.rows[0].cells[0].text = 'Report Approved by:'
        table8.rows[0].cells[1].text = 'Date:'

        newDoc.save('template.docx')
        print('Successfully generated template.docx')
    except:
        
        print('ERROR! Failed to generate docx file')

def GenManySTRs(s_Name_of_CSV, s_readDoc):
    # This function generates a .docx for each valid test in a .csv

    # print Targets
    print('Target CSV: ',s_CSV_target,'\nTarget DOCX: ',s_DOCX_target)

    #create a folder named after the current time and date
    s_FolderName = str(datetime.datetime.now()).replace(":","_")
    os.mkdir(s_FolderName)

    # Iterate through .csv file
    with open(s_Name_of_CSV, newline='') as csvfile:
        spamreader = csv.reader(csvfile, delimiter=',', quotechar='|')
        for row in spamreader:

            if (len(row)>3 and row[0]!="" and row[1]!="" and row[3]!=""):

                # Open a STR docx with A and B filled in, but not C or D
                # Fill in C with the Test Number
                # Fill in D with Results, Name, Current Date, and anomolies
                # save as a new Document under a new folder, named for the test

                AutoFiller_For_C_and_D(s_readDoc, s_FolderName+'/'+row[0]+'.docx', row[0], row[4], row[1], row[3], row[5])
    print('Many STRs Generated')
    return
                
def SetTargets(s_inputCSV_target, s_inputDOCX_target):
    # This function prompts user for .csv and .docx filepaths to use to run GenManySTRs()

    retval1 = input('Path for .csv file:\t')
    retval2 = input('Path for .docx file:\t')
    if (retval1[-4:]!='.csv'): retval1 = s_inputCSV_target
    if (retval2[-5:]!='.docx'): retval2 = s_inputDOCX_target
    retval = [retval1, retval2]
    return retval

def PathSetup():
    # 1. change working directory to match the local directory of this script
        # then return that directory as a string
    retval = str(pathlib.Path(__file__).parent.resolve())
    # print("Script is running from: ", retval)
    os.chdir(retval)
    return retval

def DetectFiles():
    # 2. scan local directory for .csv and .docx files, making a list for each #
        # also set targets on most recently scanned .csv and .docx file
    
        # declare variables
    s_CSV_target = ''
    s_DOCX_target = ''
    listCSV = []
    listDOCX = []

        # create list of all local files
    dir_list = os.listdir()

        # Scan list for local .csv and .docx files
    for localfile in dir_list:

        if (localfile[-4:]=='.csv'):
        # put .csv files in the .csv list
            listCSV.append(localfile)

        # save the most recent .csv as the .csv target
            s_CSV_target = localfile

        if (localfile[-5:]=='.docx'):
        # put .docx files in the .docx list
            listDOCX.append(localfile)

        # save the most recent .docx as the .docx target
            s_DOCX_target = localfile
    
    # 3. Check if exactly 1 .docx and exactly 1 .csv file detected
        # if they are, skip Main Menu make the generate the STRs using those
    b_SkipMain = (len(listDOCX)==1 and len(listCSV)==1)

    retval = [b_SkipMain, s_CSV_target, s_DOCX_target]
    return retval

def MainMenu( s_inputCSV_target, s_inputDOCX_target):
        
    # 4. If not skipped, open the Main Menu and prompt the user for their desired action

    s_MainMenu = '''Main Menu - input the number
    \t(1) Generate .csv template
    \t(2) Generate .docx template
    \t(3) Change targets
    \t(4) Generate STRs from targets
    \t(0) EXIT\n\t'''
    b_repeat = True
    while (b_repeat):
        print('Target CSV: ',s_inputCSV_target,'\nTarget DOCX: ',s_inputDOCX_target)
        s_MenuChoice = input(s_MainMenu)
        match s_MenuChoice:
            case '0':
                b_repeat = False
            case '1':
                GenCSV_template()
            case '2':
                GenDOCX_template()
            case '3':
                templist = SetTargets(s_inputCSV_target, s_inputDOCX_target)
                s_inputCSV_target = templist[0]
                s_inputDOCX_target = templist[1]
            case '4':
                GenManySTRs(s_inputCSV_target, s_inputDOCX_target)
            case _:
                print('invalid input')
    return

###### CODE INTRO:

# 1. Change working directory to match the local directory of this script

# 2. Scan local directory for .csv and .docx files
#    set targets on most recently scanned .csv and .docx file

# 3. Check if exactly 1 .docx and exactly 1 .csv file detected
#    if they are, skip Main Menu make the generate the STRs using those

# 4. If not skipped, open the Main Menu and prompt the user for their desired action

###### CODE START:

# 1. change working directory to match the local directory of this script 
#    getting current path and switching working directory to it
PathSetup()

# 2. scan local directory for .csv and .docx files 
#    set targets on most recently scanned .csv and .docx file
list_TargetFiles = DetectFiles()

b_SkipMainMenu = list_TargetFiles[0]
s_CSV_target = list_TargetFiles[1]
s_DOCX_target = list_TargetFiles[2]

# 3. Check if exactly 1 .docx and exactly 1 .csv file detected
#    if they are, skip Main Menu make the generate the STRs using those

if (b_SkipMainMenu):
    GenManySTRs(s_CSV_target, s_DOCX_target)
else:

# 4. If not skipped, open the Main Menu and prompt the user for their desired action
    MainMenu(s_CSV_target, s_DOCX_target)

print("END OF SCRIPT")